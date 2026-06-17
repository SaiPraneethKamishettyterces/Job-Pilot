#!/usr/bin/env python3
"""Deterministic resume renderer for the ATS resume-tailoring skill.

Takes structured resume content (see ../references/resume_content_schema.json)
and produces a file in the requested format, applying the EXACT formatting spec
in ../references/formatting_spec.md. This is what makes the layout identical on
every run — the model supplies content, this script owns typography.

Standalone: depends only on python-docx (DOCX) and reportlab (PDF), both imported
lazily. Markdown / plain text need no third-party libs.

CLI:
    python render_resume.py --content content.json --format docx --out tailored
    cat content.json | python render_resume.py --content - --format markdown --out tailored
"""

from __future__ import annotations

import argparse
import json
import sys

# ---- formatting spec constants (mirror references/formatting_spec.md) --------
DEFAULT_FONT = "Calibri"
PDF_FALLBACK_FONT = "Helvetica"  # reportlab has no bundled Calibri
SECTION_ORDER = [
    "PROFESSIONAL SUMMARY",
    "TECHNICAL SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
    "CERTIFICATIONS",
]


# ============================================================ content helpers
def _contact_line(contact: dict) -> str:
    parts = []
    if contact.get("location"):
        parts.append(contact["location"])
    if contact.get("phone"):
        parts.append(contact["phone"])
    if contact.get("email"):
        parts.append(contact["email"])
    for link in contact.get("links", []) or []:
        if link:
            parts.append(link)
    return " | ".join(parts)


def _skill_lines(technical_skills: list) -> list[str]:
    lines = []
    for cat in technical_skills or []:
        items = ", ".join(cat.get("items", []) or [])
        if cat.get("category") and items:
            lines.append(f"{cat['category']}: {items}")
    return lines


# ============================================================ Markdown / text
def to_markdown(content: dict) -> str:
    c = content.get("contact", {})
    out: list[str] = []
    out.append(f"# {c.get('name', '')}".rstrip())
    line = _contact_line(c)
    if line:
        out.append(line)
    out.append("")

    if content.get("professional_summary"):
        out += ["## PROFESSIONAL SUMMARY", content["professional_summary"], ""]

    skills = _skill_lines(content.get("technical_skills", []))
    if skills:
        out.append("## TECHNICAL SKILLS")
        out += [f"- {s}" for s in skills]
        out.append("")

    if content.get("experience"):
        out.append("## PROFESSIONAL EXPERIENCE")
        for role in content["experience"]:
            head = f"**{role.get('title', '')} | {role.get('company', '')}**"
            sub = " | ".join(p for p in (role.get("location"), role.get("dates")) if p)
            out.append(head)
            if sub:
                out.append(f"_{sub}_")
            for b in role.get("bullets", []) or []:
                out.append(f"- {b}")
            out.append("")

    if content.get("projects"):
        out.append("## PROJECTS")
        for proj in content["projects"]:
            tools = ", ".join(proj.get("tools", []) or [])
            head = f"**{proj.get('name', '')}**"
            if tools:
                head += f" | {tools}"
            out.append(head)
            for b in proj.get("bullets", []) or []:
                out.append(f"- {b}")
            out.append("")

    if content.get("education"):
        out.append("## EDUCATION")
        for ed in content["education"]:
            head = f"**{ed.get('degree', '')} | {ed.get('institution', '')}**"
            sub = " | ".join(p for p in (ed.get("location"), ed.get("dates")) if p)
            out.append(head)
            if sub:
                out.append(f"_{sub}_")
            if ed.get("details"):
                out.append(ed["details"])
            out.append("")

    if content.get("certifications"):
        out.append("## CERTIFICATIONS")
        out += [f"- {x}" for x in content["certifications"]]
        out.append("")

    return "\n".join(out).rstrip() + "\n"


def to_text(content: dict) -> str:
    """Plain text: like markdown but with underlined headings, no markdown marks."""
    md = to_markdown(content)
    lines = []
    for raw in md.splitlines():
        if raw.startswith("# "):
            name = raw[2:]
            lines += [name, "=" * len(name)]
        elif raw.startswith("## "):
            head = raw[3:]
            lines += ["", head, "-" * len(head)]
        else:
            lines.append(raw.replace("**", "").replace("_", ""))
    return "\n".join(lines).rstrip() + "\n"


# ============================================================ DOCX
def to_docx(content: dict, out_path: str, font: str = DEFAULT_FONT) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Page: US Letter, 0.5in margins all sides.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.5))

    # Base style: default font + size + single spacing, no extra space.
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(10)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    def para(space_before=0, space_after=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if align is not None:
            p.alignment = align
        return p

    def run(p, text, size, bold=False):
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold = bold
        return r

    def bottom_border(p):
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")       # ~0.75pt thin rule
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "808080")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def heading(text):
        p = para(space_before=6, space_after=2)
        run(p, text.upper(), 11, bold=True)
        bottom_border(p)

    def bullet(text, last=False):
        p = para(space_before=0, space_after=(3 if last else 0))
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run(p, f"•  {text}", 10)

    c = content.get("contact", {})

    # Name (16pt bold, centered) + contact (9.5pt centered).
    name_p = para(space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(name_p, c.get("name", ""), 16, bold=True)
    contact_text = _contact_line(c)
    if contact_text:
        cp = para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        run(cp, contact_text, 9.5)

    # Professional summary.
    if content.get("professional_summary"):
        heading("Professional Summary")
        sp = para(space_after=3)
        run(sp, content["professional_summary"], 10)

    # Technical skills.
    skills = _skill_lines(content.get("technical_skills", []))
    if skills:
        heading("Technical Skills")
        for i, line in enumerate(skills):
            cat, _, items = line.partition(": ")
            sp = para(space_after=(3 if i == len(skills) - 1 else 0))
            run(sp, f"{cat}: ", 10, bold=True)
            run(sp, items, 10)

    # Experience.
    if content.get("experience"):
        heading("Professional Experience")
        for role in content["experience"]:
            title_p = para(space_before=4)
            run(title_p, role.get("title", ""), 10.5, bold=True)
            run(title_p, f" | {role.get('company', '')}", 10.5, bold=True)
            sub = " | ".join(p for p in (role.get("location"), role.get("dates")) if p)
            if sub:
                run(para(), sub, 10)
            bullets = role.get("bullets", []) or []
            for i, b in enumerate(bullets):
                bullet(b, last=(i == len(bullets) - 1))

    # Projects.
    if content.get("projects"):
        heading("Projects")
        for proj in content["projects"]:
            head_p = para(space_before=4)
            run(head_p, proj.get("name", ""), 10.5, bold=True)
            tools = ", ".join(proj.get("tools", []) or [])
            if tools:
                run(head_p, f" | {tools}", 10)
            bullets = proj.get("bullets", []) or []
            for i, b in enumerate(bullets):
                bullet(b, last=(i == len(bullets) - 1))

    # Education.
    if content.get("education"):
        heading("Education")
        for ed in content["education"]:
            ep = para(space_before=4)
            run(ep, ed.get("degree", ""), 10.5, bold=True)
            run(ep, f" | {ed.get('institution', '')}", 10.5, bold=True)
            sub = " | ".join(p for p in (ed.get("location"), ed.get("dates")) if p)
            if sub:
                run(para(), sub, 10)
            if ed.get("details"):
                run(para(space_after=3), ed["details"], 10)

    # Certifications.
    if content.get("certifications"):
        heading("Certifications")
        certs = content["certifications"]
        for i, x in enumerate(certs):
            bullet(x, last=(i == len(certs) - 1))

    if not out_path.endswith(".docx"):
        out_path += ".docx"
    doc.save(out_path)
    return out_path


# ============================================================ PDF
def to_pdf(content: dict, out_path: str, font: str = DEFAULT_FONT) -> str:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    base = PDF_FALLBACK_FONT  # Calibri isn't bundled with reportlab
    bold = f"{base}-Bold"

    def st(name, size, leading=None, bold_=False, align=None, space_before=0, space_after=0):
        return ParagraphStyle(
            name,
            fontName=bold if bold_ else base,
            fontSize=size,
            leading=leading or size + 1.5,
            alignment=align if align is not None else 0,
            spaceBefore=space_before,
            spaceAfter=space_after,
        )

    name_style = st("name", 16, bold_=True, align=TA_CENTER)
    contact_style = st("contact", 9.5, align=TA_CENTER, space_after=4)
    head_style = st("head", 11, bold_=True, space_before=6, space_after=2)
    role_style = st("role", 10.5, bold_=True, space_before=4)
    sub_style = st("sub", 10)
    body_style = st("body", 10, leading=12)
    bullet_style = ParagraphStyle("bullet", parent=body_style, leftIndent=10, bulletIndent=0)

    if not out_path.endswith(".pdf"):
        out_path += ".pdf"
    doc = SimpleDocTemplate(
        out_path, pagesize=letter,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
    )
    flow = []
    c = content.get("contact", {})
    flow.append(Paragraph(c.get("name", ""), name_style))
    line = _contact_line(c)
    if line:
        flow.append(Paragraph(line, contact_style))

    def section(title):
        flow.append(Paragraph(title.upper(), head_style))
        flow.append(HRFlowable(width="100%", thickness=0.5, color="#808080", spaceBefore=1, spaceAfter=3))

    if content.get("professional_summary"):
        section("Professional Summary")
        flow.append(Paragraph(content["professional_summary"], body_style))

    skills = _skill_lines(content.get("technical_skills", []))
    if skills:
        section("Technical Skills")
        for s in skills:
            cat, _, items = s.partition(": ")
            flow.append(Paragraph(f"<b>{cat}:</b> {items}", body_style))

    if content.get("experience"):
        section("Professional Experience")
        for role in content["experience"]:
            flow.append(Paragraph(f"{role.get('title','')} | {role.get('company','')}", role_style))
            sub = " | ".join(p for p in (role.get("location"), role.get("dates")) if p)
            if sub:
                flow.append(Paragraph(sub, sub_style))
            for b in role.get("bullets", []) or []:
                flow.append(Paragraph(b, bullet_style, bulletText="•"))
            flow.append(Spacer(1, 3))

    if content.get("projects"):
        section("Projects")
        for proj in content["projects"]:
            tools = ", ".join(proj.get("tools", []) or [])
            head = proj.get("name", "") + (f" | {tools}" if tools else "")
            flow.append(Paragraph(head, role_style))
            for b in proj.get("bullets", []) or []:
                flow.append(Paragraph(b, bullet_style, bulletText="•"))
            flow.append(Spacer(1, 3))

    if content.get("education"):
        section("Education")
        for ed in content["education"]:
            flow.append(Paragraph(f"{ed.get('degree','')} | {ed.get('institution','')}", role_style))
            sub = " | ".join(p for p in (ed.get("location"), ed.get("dates")) if p)
            if sub:
                flow.append(Paragraph(sub, sub_style))
            if ed.get("details"):
                flow.append(Paragraph(ed["details"], body_style))

    if content.get("certifications"):
        section("Certifications")
        for x in content["certifications"]:
            flow.append(Paragraph(x, bullet_style, bulletText="•"))

    doc.build(flow)
    return out_path


# ============================================================ dispatch
def render(content: dict, fmt: str, out_base: str, font: str = DEFAULT_FONT) -> str:
    """Render `content` to `fmt`, writing to out_base(+ext). Returns the path."""
    fmt = fmt.lower()
    if fmt == "docx":
        return to_docx(content, out_base, font)
    if fmt == "pdf":
        return to_pdf(content, out_base, font)
    if fmt in ("markdown", "md"):
        path = out_base if out_base.endswith(".md") else out_base + ".md"
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(to_markdown(content))
        return path
    if fmt in ("text", "txt", "plain"):
        path = out_base if out_base.endswith(".txt") else out_base + ".txt"
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(to_text(content))
        return path
    raise ValueError(f"Unsupported format: {fmt!r} (use docx|pdf|markdown|text)")


def _main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Render structured resume content to a file.")
    ap.add_argument("--content", required=True, help="Path to content JSON, or '-' for stdin.")
    ap.add_argument("--format", default="docx", help="docx | pdf | markdown | text")
    ap.add_argument("--out", required=True, help="Output path without extension.")
    ap.add_argument("--font", default=DEFAULT_FONT)
    args = ap.parse_args(argv)

    raw = sys.stdin.read() if args.content == "-" else open(args.content, encoding="utf-8").read()
    data = json.loads(raw)
    # Accept either the full {resume, analysis} object or bare resume content.
    content = data.get("resume", data)
    path = render(content, args.format, args.out, args.font)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
